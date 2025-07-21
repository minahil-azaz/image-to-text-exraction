import os
import io
from datetime import datetime
from typing import Dict, List, Optional
import streamlit as st

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from googletrans import Translator
    TRANSLATOR_AVAILABLE = True
except ImportError:
    TRANSLATOR_AVAILABLE = False

try:
    from gtts import gTTS
    import tempfile
    TTS_AVAILABLE = True
except ImportError:
    TTS_AVAILABLE = False

class TextProcessor:
    """
    Handles text processing operations including translation, export, and text-to-speech
    """
    
    def __init__(self):
        self.translator = None
        if TRANSLATOR_AVAILABLE:
            try:
                self.translator = Translator()
            except:
                pass
        
        # Supported languages for translation
        self.translation_languages = {
            'en': 'English',
            'es': 'Spanish',
            'fr': 'French',
            'de': 'German',
            'it': 'Italian',
            'pt': 'Portuguese',
            'ru': 'Russian',
            'ja': 'Japanese',
            'ko': 'Korean',
            'zh-cn': 'Chinese (Simplified)',
            'zh-tw': 'Chinese (Traditional)',
            'ar': 'Arabic',
            'hi': 'Hindi',
            'bn': 'Bengali',
            'te': 'Telugu',
            'ta': 'Tamil',
            'mr': 'Marathi',
            'gu': 'Gujarati',
            'kn': 'Kannada',
            'ml': 'Malayalam',
            'or': 'Oriya',
            'pa': 'Punjabi',
            'ur': 'Urdu',
            'ne': 'Nepali',
            'si': 'Sinhala',
            'th': 'Thai',
            'vi': 'Vietnamese',
            'id': 'Indonesian',
            'ms': 'Malay',
            'tl': 'Filipino',
            'sw': 'Swahili',
            'zu': 'Zulu',
            'af': 'Afrikaans',
            'nl': 'Dutch',
            'sv': 'Swedish',
            'no': 'Norwegian',
            'da': 'Danish',
            'fi': 'Finnish',
            'pl': 'Polish',
            'cs': 'Czech',
            'sk': 'Slovak',
            'hu': 'Hungarian',
            'ro': 'Romanian',
            'bg': 'Bulgarian',
            'hr': 'Croatian',
            'sr': 'Serbian',
            'sl': 'Slovenian',
            'et': 'Estonian',
            'lv': 'Latvian',
            'lt': 'Lithuanian',
            'tr': 'Turkish',
            'el': 'Greek',
            'he': 'Hebrew',
            'fa': 'Persian',
            'ps': 'Pashto',
            'ku': 'Kurdish',
            'sd': 'Sindhi',
            'bal': 'Balochi',
            'ceb': 'Cebuano',
            'ilo': 'Ilocano',
            'war': 'Waray',
            'hil': 'Hiligaynon',
            'bik': 'Bikol',
            'pam': 'Kapampangan',
            'pag': 'Pangasinan',
            'km': 'Khmer',
            'lo': 'Lao',
            'my': 'Burmese',
            'dv': 'Dhivehi',
            'as': 'Assamese',
            'bho': 'Bhojpuri',
            'awa': 'Awadhi',
            'mai': 'Maithili',
            'mag': 'Magahi',
            'raj': 'Rajasthani',
            'kon': 'Konkani',
            'tcy': 'Tulu',
            'bo': 'Tibetan',
            'dz': 'Dzongkha',
            'new': 'Newari',
            'syl': 'Sylheti',
            'kha': 'Khasi',
            'gar': 'Garo',
            'mni': 'Manipuri',
            'brx': 'Bodo',
            'sat': 'Santali',
            'kui': 'Kui',
            'gon': 'Gondi',
            'kru': 'Kurukh',
            'sad': 'Sadan',
            'ho': 'Ho',
            'mwr': 'Marwari',
            'wbr': 'Wagdi',
            'bgc': 'Haryanvi',
            'hne': 'Chhattisgarhi',
            'kfy': 'Kumaoni',
            'bfy': 'Bagheli'
        }
    
    def translate_text(self, text: str, target_language: str, source_language: str = 'auto') -> Dict:
        """
        Translate text to target language
        
        Args:
            text: Text to translate
            target_language: Target language code
            source_language: Source language code (auto-detect if 'auto')
            
        Returns:
            Dictionary with translation results
        """
        if not TRANSLATOR_AVAILABLE or not self.translator:
            return {
                'success': False,
                'error': 'Translation service not available. Please install googletrans==4.0.0rc1'
            }
        
        if not text.strip():
            return {
                'success': False,
                'error': 'No text to translate'
            }
        
        try:
            # Translate the text
            translation = self.translator.translate(
                text, 
                dest=target_language, 
                src=source_language
            )
            
            return {
                'success': True,
                'original_text': text,
                'translated_text': translation.text,
                'source_language': translation.src,
                'target_language': target_language,
                'confidence': getattr(translation, 'confidence', None)
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Translation failed: {str(e)}'
            }
    
    def export_to_txt(self, text: str, filename: str = None) -> bytes:
        """
        Export text to TXT format
        
        Args:
            text: Text to export
            filename: Optional filename
            
        Returns:
            Bytes content of the TXT file
        """
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"extracted_text_{timestamp}.txt"
        
        # Add metadata header
        header = f"""# Extracted Text from Image
# Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
# Filename: {filename}

"""
        
        content = header + text
        return content.encode('utf-8')
    
    def export_to_docx(self, text: str, filename: str = None, metadata: Dict = None) -> bytes:
        """
        Export text to DOCX format
        
        Args:
            text: Text to export
            filename: Optional filename
            metadata: Additional metadata to include
            
        Returns:
            Bytes content of the DOCX file
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available. Please install it: pip install python-docx")
        
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"extracted_text_{timestamp}.docx"
        
        # Create document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Extracted Text from Image', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        if metadata:
            metadata_para = doc.add_paragraph()
            metadata_para.add_run('Metadata:').bold = True
            for key, value in metadata.items():
                metadata_para.add_run(f'\n{key}: {value}')
        
        # Add timestamp
        timestamp_para = doc.add_paragraph()
        timestamp_para.add_run(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add separator
        doc.add_paragraph('=' * 50)
        
        # Add extracted text
        text_para = doc.add_paragraph()
        text_para.add_run(text)
        
        # Save to bytes
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return docx_bytes.getvalue()
    
    def text_to_speech(self, text: str, language: str = 'en', filename: str = None) -> bytes:
        """
        Convert text to speech
        
        Args:
            text: Text to convert
            language: Language code for TTS
            filename: Optional filename
            
        Returns:
            Bytes content of the audio file
        """
        if not TTS_AVAILABLE:
            return None
        
        if not text.strip():
            return None
        
        try:
            # Create TTS object
            tts = gTTS(text=text, lang=language, slow=False)
            
            # Save to bytes
            audio_bytes = io.BytesIO()
            tts.write_to_fp(audio_bytes)
            audio_bytes.seek(0)
            
            return audio_bytes.getvalue()
            
        except Exception as e:
            st.error(f"Text-to-speech failed: {str(e)}")
            return None
    
    def clean_text(self, text: str) -> str:
        """
        Clean and format extracted text
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Remove extra whitespace
        text = ' '.join(text.split())
        
        # Fix common OCR errors
        text = text.replace('|', 'I')  # Common OCR error
        text = text.replace('0', 'O')  # Common OCR error in some fonts
        text = text.replace('1', 'l')  # Common OCR error
        
        # Capitalize first letter of sentences
        sentences = text.split('. ')
        sentences = [s.capitalize() for s in sentences if s.strip()]
        text = '. '.join(sentences)
        
        return text
    
    def extract_structured_data(self, text: str) -> Dict:
        """
        Extract structured data from text (emails, phone numbers, dates, etc.)
        
        Args:
            text: Input text
            
        Returns:
            Dictionary with extracted structured data
        """
        import re
        
        structured_data = {
            'emails': [],
            'phone_numbers': [],
            'dates': [],
            'urls': [],
            'numbers': []
        }
        
        # Extract emails
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        structured_data['emails'] = re.findall(email_pattern, text)
        
        # Extract phone numbers
        phone_pattern = r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        structured_data['phone_numbers'] = re.findall(phone_pattern, text)
        
        # Extract URLs
        url_pattern = r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
        structured_data['urls'] = re.findall(url_pattern, text)
        
        # Extract numbers
        number_pattern = r'\b\d+(?:\.\d+)?\b'
        structured_data['numbers'] = re.findall(number_pattern, text)
        
        # Extract dates (basic pattern)
        date_pattern = r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b'
        structured_data['dates'] = re.findall(date_pattern, text)
        
        return structured_data
    
    def get_translation_languages(self) -> Dict[str, str]:
        """Get available translation languages"""
        return self.translation_languages
    
    def get_language_name(self, language_code: str) -> str:
        """Get human-readable language name"""
        return self.translation_languages.get(language_code, language_code)
    
    def validate_language(self, language_code: str) -> bool:
        """Check if language code is supported"""
        return language_code in self.translation_languages
    
    def get_word_count(self, text: str) -> Dict:
        """
        Get text statistics
        
        Args:
            text: Input text
            
        Returns:
            Dictionary with text statistics
        """
        if not text:
            return {
                'characters': 0,
                'words': 0,
                'sentences': 0,
                'paragraphs': 0
            }
        
        # Count characters (excluding spaces)
        characters = len(text.replace(' ', ''))
        
        # Count words
        words = len(text.split())
        
        # Count sentences (basic)
        sentences = len([s for s in text.split('.') if s.strip()])
        
        # Count paragraphs
        paragraphs = len([p for p in text.split('\n') if p.strip()])
        
        return {
            'characters': characters,
            'words': words,
            'sentences': sentences,
            'paragraphs': paragraphs
        } 